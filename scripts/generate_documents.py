import os
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
DOCS_DIR.mkdir(exist_ok=True)

SOURCE_FILES = [
    ROOT / "AUTISM_CDSS_DOCUMENTATION.md",
    ROOT / "README.md",
]

OUTPUT_DOCX = DOCS_DIR / "Project_Documentation.docx"
OUTPUT_PDF = DOCS_DIR / "Project_Documentation.pdf"


def ensure_packages():
    try:
        import docx  # noqa: F401
        from reportlab.lib.pagesizes import A4  # noqa: F401
    except Exception:
        print("Installing required packages: python-docx reportlab")
        os.system(f"{sys.executable} -m pip install --quiet python-docx reportlab")


def read_sources():
    parts = []
    for f in SOURCE_FILES:
        if f.exists():
            parts.append(f"# Source: {f.name}\n\n")
            parts.append(f.read_text(encoding="utf-8"))
        else:
            parts.append(f"# Source: {f.name} (MISSING)\n\n")
    return "\n\n".join(parts)


def make_docx(text: str, out_path: Path):
    from docx import Document

    doc = Document()

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("• ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:])
            p.style = "List Bullet"
        else:
            doc.add_paragraph(stripped)

    doc.save(out_path)
    print(f"Wrote DOCX: {out_path}")


def make_pdf(text: str, out_path: Path):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.pagesizes import A4
    import re

    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    h1 = ParagraphStyle("Heading1", parent=styles["Heading1"]) if "Heading1" in styles else ParagraphStyle("Heading1", parent=normal, fontSize=18, leading=22)
    h2 = ParagraphStyle("Heading2", parent=styles.get("Heading2", normal), fontSize=14, leading=18)

    doc = SimpleDocTemplate(str(out_path), pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
    flow = []

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            flow.append(Spacer(1, 6))
            continue
        if stripped.startswith("# "):
            clean = re.sub(r"<[^>]+>", "", stripped[2:])
            flow.append(Paragraph(clean, h1))
        elif stripped.startswith("## "):
            clean = re.sub(r"<[^>]+>", "", stripped[3:])
            flow.append(Paragraph(clean, h2))
        elif stripped.startswith("### "):
            clean = re.sub(r"<[^>]+>", "", stripped[4:])
            try:
                flow.append(Paragraph(clean, styles["Heading3"]))
            except Exception:
                flow.append(Paragraph(clean, h2))
        elif stripped.startswith("- ") or stripped.startswith("• ") or stripped.startswith("* "):
            clean = re.sub(r"<[^>]+>", "", stripped[2:])
            flow.append(Paragraph("\u2022 " + clean, normal))
        else:
            clean = re.sub(r"<[^>]+>", "", stripped)
            flow.append(Paragraph(clean.replace("&", "&amp;"), normal))

    doc.build(flow)
    print(f"Wrote PDF: {out_path}")


def main():
    ensure_packages()
    text = read_sources()
    make_docx(text, OUTPUT_DOCX)
    make_pdf(text, OUTPUT_PDF)


if __name__ == "__main__":
    main()
